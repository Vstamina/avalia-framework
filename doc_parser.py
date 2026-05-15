# doc_parser.py

import os
import re
from typing import Dict, List, Tuple

from docx import Document
from openpyxl import load_workbook


class DocumentParser:
    """
    Parser de documentos para o Framework AVALIA.

    Lê arquivos .docx e .xlsx, limpa trechos brutos e tenta mapear evidências
    para as dimensões do framework.

    Saída esperada pelo app.py:
    {
        "mappings": [
            {
                "dimensao": "ancoragem",
                "evidencias_encontradas": [...]
            }
        ]
    }
    """

    def __init__(self, max_evidences_per_dimension: int = 6):
        self.max_evidences_per_dimension = max_evidences_per_dimension

        self.dimension_keywords = {
            "ancoragem": [
                "pós-venda",
                "pos-venda",
                "acompanhamento",
                "contrato",
                "responsável",
                "responsaveis",
                "início",
                "inicio",
                "handoff",
                "transição",
                "transicao",
                "saúde contratual",
                "saude contratual",
            ],
            "ecossistema": [
                "stakeholder",
                "ator",
                "atores",
                "decisor",
                "influenciador",
                "rh",
                "sesmt",
                "liderança",
                "lideranca",
                "trabalhador",
                "usuário final",
                "usuario final",
                "área técnica",
                "area tecnica",
            ],
            "evidencias": [
                "indicador",
                "indicadores",
                "dados",
                "evidência",
                "evidencia",
                "relatório",
                "relatorio",
                "painel",
                "dashboard",
                "métrica",
                "metrica",
                "monitoramento",
                "histórico",
                "historico",
            ],
            "jornada": [
                "jornada",
                "etapa",
                "etapas",
                "experiência",
                "experiencia",
                "ponto de contato",
                "régua de comunicação",
                "regua de comunicacao",
                "renovação",
                "renovacao",
                "comunicação",
                "comunicacao",
            ],
            "governanca": [
                "governança",
                "governanca",
                "papéis",
                "papeis",
                "responsabilidades",
                "rito",
                "ritual",
                "comitê",
                "comite",
                "fluxo",
                "escalonamento",
                "sla",
                "prazo",
                "decisão",
                "decisao",
            ],
            "acao_estrategica": [
                "ação",
                "acao",
                "plano de ação",
                "plano de acao",
                "melhoria",
                "aprendizagem",
                "resultado",
                "impacto",
                "retenção",
                "retencao",
                "expansão",
                "expansao",
                "renovação",
                "renovacao",
                "valor",
            ],
            "qualificacao": [
                "classificação",
                "classificacao",
                "categoria",
                "taxonomia",
                "criticidade",
                "causa",
                "tipo de feedback",
                "natureza",
                "manifestação",
                "manifestacao",
                "reclamação",
                "reclamacao",
                "sugestão",
                "sugestao",
                "elogio",
            ],
            "ecossistema_fb": [
                "origem do feedback",
                "emissor",
                "quem falou",
                "devolutiva",
                "interlocutor",
                "cliente",
                "trabalhador",
                "rh",
                "liderança",
                "lideranca",
                "decisor",
                "confidencialidade",
            ],
            "evidencias_fb": [
                "registro de feedback",
                "base de feedback",
                "feedback registrado",
                "recorrência",
                "recorrencia",
                "reincidência",
                "reincidencia",
                "padrão",
                "padrao",
                "histórico de feedback",
                "historico de feedback",
                "evidência do feedback",
                "evidencia do feedback",
            ],
            "jornada_fb": [
                "feedback na jornada",
                "etapa do feedback",
                "fase da jornada",
                "momento da jornada",
                "fricção",
                "friccao",
                "ponto de atrito",
                "experiência do cliente",
                "experiencia do cliente",
            ],
            "governanca_fb": [
                "tratativa",
                "fluxo de feedback",
                "prazo de resposta",
                "sla",
                "responsável pelo feedback",
                "responsavel pelo feedback",
                "escalonamento",
                "fechamento de loop",
                "loop",
                "devolutiva formal",
            ],
            "acao_fb": [
                "ação baseada em feedback",
                "acao baseada em feedback",
                "melhoria a partir do feedback",
                "aprendizado",
                "aprendizagem",
                "ação corretiva",
                "acao corretiva",
                "prevenção",
                "prevencao",
                "reincidência",
                "reincidencia",
                "efetividade",
            ],
        }

    # ======================================================
    # Função principal
    # ======================================================

    def parse_file(self, file_path: str) -> Dict:
        extension = os.path.splitext(file_path)[1].lower()

        try:
            if extension == ".docx":
                raw_segments = self._extract_docx_segments(file_path)
            elif extension == ".xlsx":
                raw_segments = self._extract_xlsx_segments(file_path)
            else:
                return {
                    "error": f"Formato não suportado: {extension}",
                    "mappings": [],
                }

            clean_segments = self._clean_segments(raw_segments)
            mappings = self._map_segments_to_dimensions(clean_segments)

            return {
                "total_raw_segments": len(raw_segments),
                "total_clean_segments": len(clean_segments),
                "mappings": mappings,
            }

        except Exception as error:
            return {
                "error": str(error),
                "mappings": [],
            }

    # ======================================================
    # Extração de Word
    # ======================================================

    def _extract_docx_segments(self, file_path: str) -> List[str]:
        document = Document(file_path)
        segments = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                segments.append(text)

        for table in document.tables:
            for row in table.rows:
                row_values = []

                for cell in row.cells:
                    value = cell.text.strip()
                    if value:
                        row_values.append(value)

                if row_values:
                    segments.append(" | ".join(row_values))

        return segments

    # ======================================================
    # Extração de Excel
    # ======================================================

    def _extract_xlsx_segments(self, file_path: str) -> List[str]:
        workbook = load_workbook(file_path, data_only=True)
        segments = []

        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values = []

                for value in row:
                    cleaned = self._normalize_cell_value(value)
                    if cleaned:
                        values.append(cleaned)

                if values:
                    segments.append(" | ".join(values))

        return segments

    def _normalize_cell_value(self, value) -> str:
        if value is None:
            return ""

        text = str(value).strip()

        if not text:
            return ""

        lower = text.lower()

        blocked_values = {
            "nan",
            "none",
            "null",
            "unnamed",
            "unnamed: 0",
            "unnamed: 1",
            "unnamed: 2",
        }

        if lower in blocked_values:
            return ""

        if lower.startswith("unnamed"):
            return ""

        return text

    # ======================================================
    # Limpeza semântica
    # ======================================================

    def _clean_segments(self, raw_segments: List[str]) -> List[str]:
        clean_segments = []
        seen = set()

        for segment in raw_segments:
            text = self._normalize_text(segment)

            if self._should_discard_segment(text):
                continue

            text = self._repair_segment(text)

            if self._should_discard_segment(text):
                continue

            key = text.lower()

            if key in seen:
                continue

            seen.add(key)
            clean_segments.append(text)

        return clean_segments

    def _normalize_text(self, text: str) -> str:
        text = str(text)
        text = text.replace("\n", " ")
        text = text.replace("\t", " ")
        text = re.sub(r"\s+", " ", text)
        text = text.strip(" -•|;:,")
        return text.strip()

    def _repair_segment(self, text: str) -> str:
        """
        Pequenos reparos para transformar trechos de tabela em frases mais legíveis.
        Não inventa conteúdo: apenas melhora a costura do que já veio no arquivo.
        """

        if " | " in text:
            parts = [part.strip() for part in text.split("|") if part.strip()]
            parts = [
                part for part in parts
                if part.lower() not in {"nan", "none", "null"}
                and not part.lower().startswith("unnamed")
            ]

            if len(parts) >= 2:
                text = "; ".join(parts)

        text = re.sub(r"\s+", " ", text).strip()

        return text

    def _should_discard_segment(self, text: str) -> bool:
        if not text:
            return True

        lower = text.lower().strip()

        blocked_exact = {
            "nan",
            "none",
            "null",
            "unnamed",
            "relatório",
            "relatorio",
            "diagnóstico",
            "diagnostico",
        }

        if lower in blocked_exact:
            return True

        if lower.startswith("unnamed"):
            return True

        if lower.endswith(":") and len(lower.split()) <= 5:
            return True

        if len(text) < 35:
            return True

        if len(text.split()) < 6:
            return True

        if self._looks_like_raw_dict(text):
            return True

        if self._has_too_many_empty_markers(text):
            return True

        return False

    def _looks_like_raw_dict(self, text: str) -> bool:
        return (
            "{" in text
            and "}" in text
            and ":" in text
            and len(text) > 80
        )

    def _has_too_many_empty_markers(self, text: str) -> bool:
        lower = text.lower()
        markers = ["nan", "none", "null", "unnamed"]
        count = sum(lower.count(marker) for marker in markers)
        return count >= 2

    # ======================================================
    # Mapeamento para dimensões
    # ======================================================

    def _map_segments_to_dimensions(self, segments: List[str]) -> List[Dict]:
        evidence_by_dimension = {
            dimension: []
            for dimension in self.dimension_keywords.keys()
        }

        for segment in segments:
            dimension_scores = self._score_segment_dimensions(segment)

            if not dimension_scores:
                continue

            best_dimension, best_score = dimension_scores[0]

            if best_score <= 0:
                continue

            if len(evidence_by_dimension[best_dimension]) >= self.max_evidences_per_dimension:
                continue

            evidence_by_dimension[best_dimension].append(segment)

        mappings = []

        for dimension, evidences in evidence_by_dimension.items():
            mappings.append(
                {
                    "dimensao": dimension,
                    "evidencias_encontradas": evidences,
                }
            )

        return mappings

    def _score_segment_dimensions(self, segment: str) -> List[Tuple[str, int]]:
        lower = segment.lower()
        scores = []

        for dimension, keywords in self.dimension_keywords.items():
            score = 0

            for keyword in keywords:
                if keyword.lower() in lower:
                    score += 1

            if score > 0:
                scores.append((dimension, score))

        scores.sort(key=lambda item: item[1], reverse=True)
        return scores