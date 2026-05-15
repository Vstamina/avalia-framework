# report_generator.py

import os
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from framework_engine import (
    AssessmentResult,
    summarize_assessment,
    get_maturity_ruler_rows,
)


# ==========================================================
# FORMATAÇÃO BÁSICA
# ==========================================================

def _format_run(run, bold=False, size=11):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(6)

    for run in paragraph.runs:
        _format_run(run)


def _add_paragraph(doc, text: str, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _format_run(run, bold=bold)
    _format_paragraph(paragraph, alignment=alignment)
    return paragraph


def _set_cell_text(cell, text, bold=False):
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(str(text))
    _format_run(run, bold=bold)

    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)


def _add_table(doc, headers, data_rows):
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for col_idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[col_idx], header, bold=True)

    for row_idx, row in enumerate(data_rows):
        for col_idx, value in enumerate(row):
            _set_cell_text(table.rows[row_idx + 1].cells[col_idx], value)

    doc.add_paragraph()
    return table


def _configure_document_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _format_heading(heading):
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for run in heading.runs:
        _format_run(run, bold=True)


# ==========================================================
# HELPERS VISUAIS PARA ONE-PAGE EXECUTIVO
# ==========================================================

def _format_number_br(value):
    try:
        return f"{float(value):.2f}".replace(".", ",")
    except (TypeError, ValueError):
        return str(value)


def _shade_cell(cell, fill_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color.replace("#", ""))
    tc_pr.append(shd)


def _set_cell_vertical_center(cell):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_text_visual(
    cell,
    text,
    bold=False,
    size=10,
    color="1F2937",
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    _set_cell_vertical_center(cell)


def _set_table_borders_light(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color.replace("#", ""))


def _add_kpi_card(table, col_idx, label, value, help_text, accent_color="0FA3B1"):
    cell = table.rows[0].cells[col_idx]
    _shade_cell(cell, "FFFFFF")
    _set_cell_vertical_center(cell)

    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    r1.font.name = "Arial"
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor.from_string("6B7280")
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(str(value))
    r2.font.name = "Arial"
    r2.font.size = Pt(18)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(accent_color)
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(help_text)
    r3.font.name = "Arial"
    r3.font.size = Pt(7)
    r3.font.color.rgb = RGBColor.from_string("6B7280")
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


# ==========================================================
# FUNÇÕES DE APOIO AO RELATÓRIO
# ==========================================================

def _safe_join(items: List[str], fallback: str = "não identificada"):
    if not items:
        return fallback

    return ", ".join(items)


def _get_results_by_layer(results: List[AssessmentResult], layer: int):
    return [item for item in results if item.camada == layer]


def _get_critical_dimensions_text(summary):
    if not summary.dimensoes_criticas:
        return "Não foram identificadas dimensões críticas no patamar Incipiente ou Inicial."

    return ", ".join(summary.dimensoes_criticas)


def _get_priority_focus(summary):
    if summary.dimensoes_criticas:
        return "Governança, jornada e feedback qualificado"

    if summary.dimensoes_em_evolucao:
        return "Integração, evidências e melhoria contínua"

    return "Manutenção da maturidade e inovação contínua"


def _add_lacunas_table(doc, item: AssessmentResult):
    rows = []

    for lacuna in item.lacunas:
        rows.append([
            lacuna.lacuna,
            lacuna.evidencia,
            lacuna.impacto,
        ])

    if not rows:
        rows = [[
            "Lacuna não especificada",
            "Evidência não informada",
            "Impacto a validar",
        ]]

    _add_table(
        doc,
        ["Lacuna", "Evidência", "Impacto gerencial"],
        rows,
    )


# ==========================================================
# ONE-PAGE EXECUTIVO
# ==========================================================

def _add_executive_one_page(doc, client_data: dict, results: List[AssessmentResult], summary):
    # Faixa superior
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True

    left = header_table.rows[0].cells[0]
    right = header_table.rows[0].cells[1]

    _shade_cell(left, "082B63")
    _shade_cell(right, "082B63")

    _set_cell_text_visual(
        left,
        "Resumo Executivo do Diagnóstico AVALIA",
        bold=True,
        size=22,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _set_cell_text_visual(
        right,
        "Logo",
        bold=True,
        size=16,
        color="082B63",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _shade_cell(right, "FFFFFF")

    doc.add_paragraph()

    # Texto introdutório
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.space_after = Pt(6)

    run = intro.add_run(
        "Esta página apresenta uma síntese executiva do diagnóstico realizado por meio do Framework AVALIA, "
        "com foco na maturidade da gestão de pós-venda e do feedback qualificado. O objetivo é oferecer uma "
        "leitura rápida dos principais indicadores, forças, lacunas e prioridades de intervenção."
    )
    _format_run(run, size=9)

    # Bloco de identificação
    info_table = doc.add_table(rows=2, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = True
    _set_table_borders_light(info_table)

    labels = [
        ("Organização", client_data.get("client_name", "Cliente")),
        ("Unidade", client_data.get("unit", "[preencher]")),
        ("Período", client_data.get("period", "")),
        ("Responsável", client_data.get("responsible", "")),
    ]

    for idx, (label, value) in enumerate(labels):
        _shade_cell(info_table.rows[0].cells[idx], "D9F3F6")
        _set_cell_text_visual(
            info_table.rows[0].cells[idx],
            label,
            bold=True,
            size=8,
            color="082B63",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        _set_cell_text_visual(
            info_table.rows[1].cells[idx],
            value,
            bold=False,
            size=9,
            color="1F2937",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    doc.add_paragraph()

    # KPIs principais
    prioridades_altas = len([
        item for item in results
        if item.prioridade_intervencao == "Alta"
    ])

    dimensoes_avaliadas = len(results)

    kpi_table = doc.add_table(rows=1, cols=5)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.autofit = True
    _set_table_borders_light(kpi_table)

    _add_kpi_card(
        kpi_table,
        0,
        "Índice geral",
        _format_number_br(summary.indice_geral),
        summary.nivel_geral,
        "0FA3B1",
    )
    _add_kpi_card(
        kpi_table,
        1,
        "Camada 1",
        _format_number_br(summary.indice_camada_1),
        "Pós-venda",
        "0B7285",
    )
    _add_kpi_card(
        kpi_table,
        2,
        "Camada 2",
        _format_number_br(summary.indice_camada_2),
        "Feedback",
        "59C3C3",
    )
    _add_kpi_card(
        kpi_table,
        3,
        "Prioridades altas",
        prioridades_altas,
        "Foco imediato",
        "E76F51",
    )
    _add_kpi_card(
        kpi_table,
        4,
        "Dimensões",
        dimensoes_avaliadas,
        "Avaliadas",
        "2E8B57",
    )

    doc.add_paragraph()

    # Visão geral do diagnóstico
    overview_title = doc.add_paragraph()
    overview_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = overview_title.add_run("Visão Geral do Diagnóstico")
    _format_run(r, bold=True, size=14)

    overview = doc.add_paragraph()
    overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    overview_text = (
        f"O diagnóstico indica índice geral de maturidade de {_format_number_br(summary.indice_geral)}, "
        f"situando a organização no patamar {summary.nivel_geral}. A Camada 1 apresentou índice "
        f"{_format_number_br(summary.indice_camada_1)} ({summary.nivel_camada_1}), enquanto a Camada 2 "
        f"apresentou índice {_format_number_br(summary.indice_camada_2)} ({summary.nivel_camada_2}). "
        "A leitura executiva sugere que a intervenção deve priorizar dimensões com baixa maturidade "
        "e alto impacto estratégico, especialmente aquelas relacionadas à governança, jornada, evidências "
        "e fechamento do ciclo de feedback."
    )

    r = overview.add_run(overview_text)
    _format_run(r, size=9)

    # Achados principais em duas colunas
    findings_table = doc.add_table(rows=3, cols=2)
    findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    findings_table.autofit = True
    _set_table_borders_light(findings_table)

    _shade_cell(findings_table.rows[0].cells[0], "082B63")
    _shade_cell(findings_table.rows[0].cells[1], "082B63")

    _set_cell_text_visual(
        findings_table.rows[0].cells[0],
        "Forças identificadas",
        bold=True,
        size=11,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text_visual(
        findings_table.rows[0].cells[1],
        "Prioridades de intervenção",
        bold=True,
        size=11,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for cell in findings_table.rows[1].cells:
        _shade_cell(cell, "EAFBFD")

    strengths = summary.dimensoes_maior_maturidade[:5]
    priorities = [
        item.dimensao for item in results
        if item.prioridade_intervencao == "Alta"
    ][:5]

    if not strengths:
        strengths = ["Não identificadas"]
    if not priorities:
        priorities = ["Não foram identificadas prioridades altas"]

    strengths_text = "\n".join([f"• {item}" for item in strengths])
    priorities_text = "\n".join([f"• {item}" for item in priorities])

    _set_cell_text_visual(
        findings_table.rows[1].cells[0],
        strengths_text,
        size=9,
        color="1F2937",
    )
    _set_cell_text_visual(
        findings_table.rows[1].cells[1],
        priorities_text,
        size=9,
        color="1F2937",
    )

    _shade_cell(findings_table.rows[2].cells[0], "FFFFFF")
    _shade_cell(findings_table.rows[2].cells[1], "FFFFFF")

    _set_cell_text_visual(
        findings_table.rows[2].cells[0],
        f"Principal força relativa: {summary.principal_forca}",
        bold=False,
        size=8,
        color="6B7280",
    )
    _set_cell_text_visual(
        findings_table.rows[2].cells[1],
        f"Principal lacuna relativa: {summary.principal_lacuna}",
        bold=False,
        size=8,
        color="6B7280",
    )

    doc.add_paragraph()

    # Roadmap executivo
    roadmap_title = doc.add_paragraph()
    roadmap_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = roadmap_title.add_run("Roadmap Executivo de Intervenção")
    _format_run(r, bold=True, size=14)

    roadmap_table = doc.add_table(rows=6, cols=3)
    roadmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roadmap_table.autofit = True
    _set_table_borders_light(roadmap_table)

    headers = ["Fase", "Horizonte", "Entregável-chave"]

    for idx, header in enumerate(headers):
        _shade_cell(roadmap_table.rows[0].cells[idx], "082B63")
        _set_cell_text_visual(
            roadmap_table.rows[0].cells[idx],
            header,
            bold=True,
            size=9,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    roadmap_rows = [
        ["Fase 1", "0-30 dias", "Alinhamento conceitual e matriz de atores"],
        ["Fase 2", "31-60 dias", "Fluxos, taxonomia e instrumentos"],
        ["Fase 3", "61-120 dias", "Piloto controlado"],
        ["Fase 4", "121-180 dias", "Avaliação, ajustes e governança"],
        ["Fase 5", "181-360 dias", "Escala e institucionalização"],
    ]

    for row_idx, row_data in enumerate(roadmap_rows, start=1):
        for col_idx, value in enumerate(row_data):
            if row_idx % 2 == 0:
                _shade_cell(roadmap_table.rows[row_idx].cells[col_idx], "F5F7FA")
            else:
                _shade_cell(roadmap_table.rows[row_idx].cells[col_idx], "FFFFFF")

            _set_cell_text_visual(
                roadmap_table.rows[row_idx].cells[col_idx],
                value,
                size=8,
                color="1F2937",
                alignment=WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT,
            )

    # Fechamento da página
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(8)

    r = footer.add_run(
        "A página seguinte apresenta o relatório analítico completo, com método, diagnóstico por dimensão, lacunas, recomendações e indicadores."
    )
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("6B7280")
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    doc.add_page_break()


# ==========================================================
# CAPA / IDENTIFICAÇÃO
# ==========================================================

def _add_identification(doc, client_data: dict):
    title = doc.add_heading(
        "RELATÓRIO DE DIAGNÓSTICO E INTERVENÇÃO ORGANIZACIONAL",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(
        "Framework AVALIA aplicado à Gestão de Pós-Venda e Feedback Qualificado em Serviços de Saúde",
        level=1,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_table(doc, ["Campo", "Conteúdo"], [
        ("Organização", client_data.get("client_name", "SESI")),
        ("Área/unidade analisada", client_data.get("unit", "[preencher]")),
        (
            "Serviços contemplados",
            client_data.get(
                "services",
                "Saúde ocupacional, segurança do trabalho, promoção da saúde",
            ),
        ),
        ("Período de coleta", client_data.get("period", datetime.now().strftime("%m/%Y"))),
        ("Responsáveis pela análise", client_data.get("responsible", "Consultoria Estratégica")),
        ("Versão do relatório", "1.0 - Draft"),
    ])

    doc.add_page_break()


# ==========================================================
# 1. SUMÁRIO EXECUTIVO
# ==========================================================

def _add_executive_summary(doc, summary):
    doc.add_heading("1. SUMÁRIO EXECUTIVO", level=1)

    _add_paragraph(
        doc,
        (
            f"O diagnóstico realizado por meio do Framework AVALIA identificou índice geral de "
            f"maturidade de {summary.indice_geral:.2f}, em escala de 0 a 3, situando o modelo "
            f"no patamar {summary.nivel_geral}. A Camada 1, voltada à gestão de pós-venda, "
            f"apresentou índice de {summary.indice_camada_1:.2f} ({summary.nivel_camada_1}), "
            f"enquanto a Camada 2, dedicada ao feedback qualificado, apresentou índice de "
            f"{summary.indice_camada_2:.2f} ({summary.nivel_camada_2})."
        ),
    )

    _add_paragraph(
        doc,
        (
            "Os achados indicam que a organização dispõe de capacidades relevantes para estruturar "
            "um modelo mais robusto de acompanhamento contratual, demonstração de valor e aprendizagem. "
            "Entretanto, essas capacidades ainda precisam ser integradas em uma lógica contínua de "
            "governança, leitura da jornada, qualificação de feedbacks e ação estratégica."
        ),
    )

    _add_paragraph(
        doc,
        (
            "A intervenção proposta recomenda a implantação gradual de um modelo integrado de pós-venda "
            "e feedback qualificado, capaz de transformar dados, manifestações e interações com clientes "
            "em evidências gerenciais, decisões, ações corretivas, oportunidades de melhoria e aprendizagem "
            "organizacional."
        ),
    )

    _add_table(doc, ["Campo", "Resultado"], [
        ("Grau geral de maturidade", summary.nivel_geral),
        ("Índice geral", f"{summary.indice_geral:.2f}"),
        ("Índice Camada 1 — Pós-Venda", f"{summary.indice_camada_1:.2f} ({summary.nivel_camada_1})"),
        ("Índice Camada 2 — Feedback Qualificado", f"{summary.indice_camada_2:.2f} ({summary.nivel_camada_2})"),
        ("Principal força identificada", summary.principal_forca),
        ("Principal lacuna identificada", summary.principal_lacuna),
        ("Dimensões críticas", _get_critical_dimensions_text(summary)),
        ("Prioridade de intervenção", _get_priority_focus(summary)),
        ("Horizonte recomendado", "Médio prazo (6-12 meses)"),
        ("Entregáveis prioritários", "Framework / Protocolo / Roadmap / Matriz de Indicadores"),
    ])

    doc.add_page_break()


# ==========================================================
# 2. CONTEXTUALIZAÇÃO
# ==========================================================

def _add_context(doc):
    doc.add_heading("2. CONTEXTUALIZAÇÃO ORGANIZACIONAL", level=1)

    _add_paragraph(
        doc,
        (
            "A organização analisada atua em contexto de alta complexidade, oferecendo serviços "
            "voltados à saúde, segurança, bem-estar e produtividade. Sua atuação envolve empresas "
            "contratantes, trabalhadores, lideranças, RH, Serviço Especializado em Engenharia de "
            "Segurança e em Medicina do Trabalho (SESMT), profissionais técnicos e áreas internas "
            "de gestão, atendimento e operação."
        ),
    )

    _add_table(doc, ["Característica", "Implicação para o pós-venda"], [
        ("Contratos contínuos", "Exigem acompanhamento durante toda a vigência."),
        ("Serviços técnicos e sensíveis", "Demandam clareza, confiança e proteção de dados."),
        ("Múltiplos atores", "Exigem diferenciação entre contratante, decisor, influenciador e usuário final."),
        ("Valor nem sempre visível", "Exigem demonstração recorrente de benefícios e resultados."),
        ("Dependência de adesão", "Exigem comunicação adequada com trabalhadores e lideranças."),
        ("Relação B2B com usuário final", "Exige gestão simultânea da empresa contratante e da experiência do trabalhador."),
    ])


# ==========================================================
# 3. MÉTODO
# ==========================================================

def _add_method(doc):
    doc.add_heading("3. MÉTODO DE DIAGNÓSTICO APLICADO", level=1)

    _add_paragraph(
        doc,
        (
            "O diagnóstico foi orientado pelo Framework AVALIA, estruturado em duas camadas analíticas. "
            "A primeira camada examina a gestão de pós-venda, considerando acompanhamento contratual, "
            "jornada, governança, dados, valor e ação estratégica. A segunda camada examina o feedback "
            "qualificado, considerando captação, registro, classificação, tratativa, devolutiva, "
            "aprendizagem e retroalimentação da gestão."
        ),
    )

    _add_table(doc, ["Camada", "Finalidade"], [
        (
            "Camada 1 — Gestão de Pós-Venda",
            "Avaliar acompanhamento de contratos, jornada, dados, governança, valor e ação estratégica após a venda.",
        ),
        (
            "Camada 2 — Feedback Qualificado",
            "Avaliar captação, registro, classificação, tratativa e transformação de feedbacks em melhoria e decisão.",
        ),
    ])

    doc.add_heading("3.1 Régua de interpretação da maturidade", level=2)

    _add_paragraph(
        doc,
        (
            "Os índices de maturidade são calculados em escala de 0 a 3. Cada dimensão recebe uma nota "
            "a partir das microperguntas ou evidências analisadas. O índice geral corresponde à média "
            "das dimensões avaliadas, podendo também ser observado por camada."
        ),
    )

    _add_table(
        doc,
        ["Intervalo", "Nível de maturidade", "Interpretação"],
        get_maturity_ruler_rows(),
    )

    _add_paragraph(
        doc,
        (
            "Índice da dimensão = soma das pontuações das microperguntas da dimensão dividida pelo "
            "número de microperguntas respondidas. Índice geral = média dos índices das dimensões "
            "avaliadas. Quando uma informação não é aplicável ao caso, recomenda-se excluí-la do "
            "denominador; quando deveria existir e não há evidência, a pontuação deve refletir ausência "
            "ou fragilidade da prática."
        ),
    )

    doc.add_page_break()


# ==========================================================
# 4. DIAGNÓSTICO POR CAMADAS
# ==========================================================

def _render_dimension(doc, item: AssessmentResult):
    doc.add_heading(f"Dimensão {item.dimensao}", level=3)

    for block in item.texto_diagnostico.split("\n\n"):
        if block.strip():
            _add_paragraph(doc, block.strip())

    _add_paragraph(
        doc,
        (
            f"Índice de maturidade: {item.score_final:.2f} "
            f"({item.nivel_maturidade}) | Prioridade: {item.prioridade_intervencao}"
        ),
        bold=True,
    )

    _add_paragraph(doc, "Lacunas identificadas:", bold=True)
    _add_lacunas_table(doc, item)

    _add_paragraph(
        doc,
        f"Recomendação preliminar: {item.recomendacao_estrategica}",
        bold=False,
    )


def _add_diagnostic_by_layers(doc, results: List[AssessmentResult]):
    doc.add_heading("4. DIAGNÓSTICO POR CAMADAS E DIMENSÕES", level=1)

    camada_1_results = _get_results_by_layer(results, 1)
    camada_2_results = _get_results_by_layer(results, 2)

    doc.add_heading("CAMADA 1 — Gestão de Pós-Venda para Serviços de Saúde", level=2)
    for item in camada_1_results:
        _render_dimension(doc, item)

    doc.add_page_break()

    doc.add_heading("CAMADA 2 — Protocolo de Feedback Qualificado", level=2)
    for item in camada_2_results:
        _render_dimension(doc, item)

    doc.add_page_break()


# ==========================================================
# 5. SÍNTESE
# ==========================================================

def _add_synthesis(doc, results: List[AssessmentResult], summary):
    doc.add_heading("5. SÍNTESE DOS ACHADOS CRÍTICOS", level=1)

    rows = []

    for item in results:
        principal_lacuna = item.lacunas[0].lacuna if item.lacunas else "Não especificada"

        rows.append([
            item.dimensao,
            f"{item.score_final:.2f}".replace(".", ","),
            item.nivel_maturidade,
            item.impacto_estrategico_aplicado,
            principal_lacuna,
            item.prioridade_intervencao,
        ])

    _add_table(
        doc,
        [
            "Dimensão",
            "Índice",
            "Maturidade",
            "Impacto estratégico",
            "Principal lacuna",
            "Prioridade",
        ],
        rows,
    )

    maior_maturidade = _safe_join(
        summary.dimensoes_maior_maturidade,
        "nenhuma dimensão destacada",
    )

    criticas = _safe_join(
        summary.dimensoes_criticas,
        "nenhuma dimensão crítica",
    )

    evolucao = _safe_join(
        summary.dimensoes_em_evolucao,
        "nenhuma dimensão no patamar Organizado",
    )

    _add_paragraph(
        doc,
        (
            f"As dimensões com maior maturidade relativa foram {maior_maturidade}. "
            f"As dimensões críticas, situadas nos patamares Incipiente ou Inicial, foram {criticas}. "
            f"As dimensões em evolução, situadas no patamar Organizado, foram {evolucao}."
        ),
    )

    _add_paragraph(
        doc,
        (
            "A leitura transversal indica que a intervenção deve priorizar os elementos que conectam "
            "governança, jornada, evidências e feedback. A maturidade do pós-venda não depende apenas "
            "da existência de contatos ou relatórios, mas da capacidade de transformar esses elementos "
            "em gestão contínua de valor, prevenção de riscos e aprendizagem organizacional."
        ),
    )

    doc.add_page_break()


# ==========================================================
# 6. INTERVENÇÃO PROPOSTA
# ==========================================================

def _add_intervention(doc):
    doc.add_heading("6. INTERVENÇÃO PROPOSTA", level=1)

    doc.add_heading("6.1 Objetivo geral da intervenção", level=2)
    _add_paragraph(
        doc,
        (
            "Estruturar um modelo integrado de pós-venda e feedback qualificado para os serviços "
            "de saúde, capaz de acompanhar contratos de forma contínua, demonstrar valor ao cliente, "
            "antecipar riscos, qualificar manifestações, orientar decisões e promover melhoria contínua."
        ),
    )

    doc.add_heading("6.2 Objetivos específicos", level=2)

    objectives = [
        "Definir o conceito operacional de pós-venda em serviços de saúde.",
        "Mapear atores, papéis e responsabilidades no ecossistema contratual.",
        "Estruturar uma jornada gerencial de pós-venda com marcos de acompanhamento.",
        "Criar critérios para identificação de contrato saudável e contrato em risco.",
        "Padronizar a captação, classificação e tratativa de feedbacks.",
        "Estabelecer governança mínima de papéis, prazos, escalonamento e devolutiva.",
        "Definir indicadores para monitorar valor percebido, resposta, reincidência e melhoria.",
        "Criar mecanismos de aprendizagem e retroalimentação do pós-venda.",
    ]

    for objective in objectives:
        doc.add_paragraph(objective, style="List Bullet")

    doc.add_heading("6.3 Artefato 1 — Framework de Gestão de Pós-Venda", level=2)
    _add_table(doc, ["Componente", "Descrição"], [
        ("Conceito de pós-venda", "Definição institucional e operacional."),
        ("Jornada de pós-venda", "Etapas, marcos, contatos, indicadores e responsáveis."),
        ("Matriz de atores", "Decisor, RH, SESMT, trabalhador, liderança e áreas internas."),
        ("Sinais de saúde contratual", "Indicadores de contrato saudável, risco e oportunidade."),
        ("Demonstração de valor", "Relatórios, reuniões, indicadores e narrativa de impacto."),
        ("Governança", "Papéis, ritos, escalonamento e tomada de decisão."),
        ("Aprendizagem", "Registro e disseminação de melhorias."),
    ])

    doc.add_heading("6.4 Artefato 2 — Protocolo de Feedback Qualificado", level=2)
    _add_table(doc, ["Componente", "Descrição"], [
        ("Conceito de feedback", "Definição do que deve ser registrado como feedback."),
        ("Taxonomia", "Reclamação, dúvida, sugestão, elogio, risco, oportunidade e demanda."),
        ("Campos mínimos", "Origem, canal, serviço, etapa, criticidade, impacto e responsável."),
        ("Criticidade", "Baixa, média, alta e crítica."),
        ("Fluxo de tratativa", "Registro, classificação, priorização, ação, resposta e fechamento."),
        ("Devolutiva", "Técnica, relacional ou executiva."),
        ("Indicadores", "Abertos, tratados, encerrados, reincidentes e tempo de resposta."),
        ("Aprendizagem", "Causa, melhoria, prevenção e retroalimentação."),
    ])

    doc.add_page_break()


# ==========================================================
# 7. ROADMAP E INDICADORES
# ==========================================================

def _add_roadmap_and_indicators(doc):
    doc.add_heading("7. ROADMAP DE AÇÃO E INDICADORES", level=1)

    fases = [
        (
            "Fase 1 — Preparação e alinhamento conceitual (0 a 30 dias)",
            [
                ("Validar conceito de pós-venda", "Gestão", "Definição aprovada", "Conceito formalizado"),
                ("Mapear atores do contrato", "Comercial + operação", "Matriz de atores", "Contratos mapeados"),
                ("Validar taxonomia inicial de feedback", "Qualidade + atendimento", "Taxonomia versão 1", "Categorias aprovadas"),
            ],
        ),
        (
            "Fase 2 — Desenho dos fluxos e instrumentos (31 a 60 dias)",
            [
                ("Desenhar jornada gerencial de pós-venda", "Operação + relacionamento", "Jornada-mãe", "Jornada validada"),
                ("Definir campos mínimos de feedback", "Qualidade + BI", "Formulário/protocolo", "Campos implementados"),
                ("Definir fluxo de escalonamento", "Gestão + áreas técnicas", "Fluxo aprovado", "SLA definido"),
            ],
        ),
        (
            "Fase 3 — Piloto controlado (61 a 120 dias)",
            [
                ("Aplicar framework em contratos-piloto", "Equipe designada", "Diagnóstico por contrato", "Número de contratos monitorados"),
                ("Registrar feedbacks conforme protocolo", "Atendimento + operação", "Base inicial", "Percentual de feedbacks classificados"),
                ("Testar devolutiva ao cliente", "Relacionamento", "Fechamento de loop", "Percentual de feedbacks com retorno"),
            ],
        ),
        (
            "Fase 4 — Avaliação e ajustes (121 a 180 dias)",
            [
                ("Avaliar resultados do piloto", "Gestão + BI", "Relatório de avaliação", "Indicadores comparados"),
                ("Ajustar taxonomia e fluxo", "Qualidade + operação", "Versão 2 do protocolo", "Ajustes implementados"),
                ("Validar expansão", "Gestão", "Plano de escala", "Decisão de expansão"),
            ],
        ),
        (
            "Fase 5 — Escala e institucionalização (181 a 360 dias)",
            [
                ("Expandir para outros serviços", "Gestão", "Plano de expansão", "Número de serviços incluídos"),
                ("Integrar dados e relatórios", "BI + áreas", "Painel de gestão", "Painel ativo"),
                ("Criar rotina de governança", "Gestão", "Comitê/ritual", "Reuniões realizadas"),
                ("Incorporar aprendizado", "Gestão + qualidade", "Base de melhorias", "Melhorias replicadas"),
            ],
        ),
    ]

    for titulo, acoes in fases:
        doc.add_heading(titulo, level=2)
        _add_table(doc, ["Ação", "Responsável", "Produto", "Indicador"], acoes)

    doc.add_heading("Indicadores recomendados", level=2)
    _add_table(doc, ["Indicador", "Finalidade"], [
        ("Percentual de contratos com jornada mapeada", "Medir estruturação do acompanhamento."),
        ("Percentual de contratos com responsável definido", "Medir governança."),
        ("Percentual de contratos com reunião intermediária", "Medir acompanhamento ativo."),
        ("Percentual de contratos com demonstração de valor", "Medir comunicação de resultados."),
        ("Taxa de renovação", "Medir continuidade."),
        ("Taxa de expansão de serviços", "Medir valor e oportunidade."),
        ("Índice de contratos em risco", "Medir capacidade preventiva."),
        ("Número de feedbacks registrados", "Medir captação."),
        ("Percentual de feedbacks classificados", "Medir qualidade do registro."),
        ("Tempo médio de primeira resposta", "Medir agilidade."),
        ("Tempo médio de resolução", "Medir eficiência."),
        ("Percentual de feedbacks encerrados com devolutiva", "Medir fechamento de loop."),
        ("Percentual de reincidência", "Medir efetividade da solução."),
        ("CSAT pós-tratativa", "Medir satisfação após resposta."),
    ])

    doc.add_page_break()


# ==========================================================
# 8. CONCLUSÃO
# ==========================================================

def _add_conclusion(doc, summary):
    doc.add_heading("8. CONCLUSÃO", level=1)

    _add_paragraph(
        doc,
        (
            f"O diagnóstico realizado por meio do Framework AVALIA demonstra que a organização se encontra "
            f"no patamar {summary.nivel_geral}, com índice geral de {summary.indice_geral:.2f}. "
            "Esse resultado indica que há capacidades a serem reconhecidas, mas também lacunas que exigem "
            "intervenção estruturada para consolidar governança, jornada, evidências e feedback qualificado."
        ),
    )

    _add_paragraph(
        doc,
        (
            "A contribuição central do Framework AVALIA é organizar a leitura do pós-venda não como uma ação "
            "isolada após a venda, mas como uma capacidade organizacional de acompanhamento, demonstração de "
            "valor, prevenção de riscos e aprendizagem. Ao integrar a gestão contratual ao tratamento qualificado "
            "dos feedbacks, o modelo permite transformar manifestações, dados e interações em decisões e melhorias."
        ),
    )

    _add_paragraph(
        doc,
        (
            "A implantação deve ocorrer de forma gradual, iniciando pelo alinhamento conceitual e pela definição "
            "de papéis, avançando para instrumentos, piloto controlado, avaliação e institucionalização. Com isso, "
            "espera-se fortalecer a capacidade da organização de acompanhar contratos, responder melhor aos clientes, "
            "ampliar retenção, qualificar a experiência e transformar feedbacks em melhoria contínua."
        ),
    )


# ==========================================================
# FUNÇÃO PRINCIPAL
# ==========================================================

def generate_report(client_data: dict, results: List[AssessmentResult], output_dir: str = "relatorios") -> str:
    doc = Document()
    _configure_document_style(doc)

    summary = summarize_assessment(results)

    _add_executive_one_page(doc, client_data, results, summary)
    _add_identification(doc, client_data)
    _add_executive_summary(doc, summary)
    _add_context(doc)
    _add_method(doc)
    _add_diagnostic_by_layers(doc, results)
    _add_synthesis(doc, results, summary)
    _add_intervention(doc)
    _add_roadmap_and_indicators(doc)
    _add_conclusion(doc, summary)

    os.makedirs(output_dir, exist_ok=True)

    client_name = client_data.get("client_name", "Cliente")
    safe_client_name = (
        client_name
        .replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    filename = f"Relatorio_Intervencao_{safe_client_name}.docx"
    path = os.path.join(output_dir, filename)

    doc.save(path)
    return path