# doc_parser.py
import pandas as pd
from docx import Document
from typing import List, Dict, Optional
import os
import gc

class DocumentParser:
    def __init__(self):
        self.dimensions_c1 = ["ancoragem", "ecossistema", "evidencias", "jornada", "governanca", "acao_estrategica"]
        self.dimensions_c2 = ["qualificacao", "ecossistema_fb", "evidencias_fb", "jornada_fb", "governanca_fb", "acao_fb"]
        
    def extract_from_docx(self, file_path):
        if not os.path.exists(file_path): return {"error": "Arquivo não encontrado"}
        try:
            doc = Document(file_path)
            extracted = {"source": os.path.basename(file_path), "type": "docx", "paragraphs": [], "tables": []}
            for para in doc.paragraphs:
                if para.text.strip(): extracted["paragraphs"].append(para.text.strip())
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data): table_data.append(row_data)
                if table_data: extracted["tables"].append(table_data)
            return extracted
        except Exception as e:
            return {"error": str(e)}
    
    def extract_from_xlsx(self, file_path, sheet_name=None):
        if not os.path.exists(file_path): return {"error": "Arquivo não encontrado"}
        xls = None
        try:
            # Abre explicitamente para controle
            xls = pd.ExcelFile(file_path)
            sheets = [sheet_name] if sheet_name else xls.sheet_names
            
            extracted_sheets = {}
            for sheet in sheets:
                df = pd.read_excel(xls, sheet_name=sheet)
                df = df.dropna(how='all')
                extracted_sheets[sheet] = df.to_dict(orient='records')
            
            return {"source": os.path.basename(file_path), "type": "xlsx", "sheets": extracted_sheets}
        except Exception as e:
            return {"error": str(e)}
        finally:
            # Garante que o arquivo é fechado antes de retornar
            if xls:
                xls.close()
            gc.collect() # Limpa memória
    
    def map_to_dimensions(self, extracted_data):
        mappings = []
        text_content = []
        if extracted_data.get("paragraphs"): text_content.extend(extracted_data["paragraphs"])
        if extracted_data.get("tables"):
            for table in extracted_data["tables"]:
                for row in table: text_content.append(" | ".join(row))
        if extracted_data.get("sheets"):
            for sheet, rows in extracted_data["sheets"].items():
                for row in rows: text_content.append(str(row))
        
        keywords_map = {
            "ancoragem": ["pós-venda", "acompanhamento", "contrato", "renovação", "relacionamento"],
            "ecossistema": ["rh", "sesmt", "trabalhador", "decisor", "stakeholder"],
            "evidencias": ["indicador", "dado", "relatório", "nps", "csat", "métrica"],
            "jornada": ["etapa", "fase", "onboarding", "ativação", "implantação"],
            "governanca": ["responsável", "papel", "ritual", "reunião", "escalona", "decisão"],
            "acao_estrategica": ["ação", "plano", "melhoria", "aprendizado", "inovação"],
            "qualificacao": ["reclamação", "dúvida", "sugestão", "elogio", "feedback", "classificação"],
            "ecossistema_fb": ["origem", "quem falou", "rh", "decisor", "canal"],
            "evidencias_fb": ["registro", "base", "histórico", "recorrência"],
            "jornada_fb": ["etapa do feedback", "momento", "fase da jornada"],
            "governanca_fb": ["prazo", "responsável feedback", "escalona", "tratativa"],
            "acao_fb": ["retorno", "fechamento", "loop", "ação tomada", "devolutiva"]
        }
        
        for dim, keywords in keywords_map.items():
            matches = []
            for text in text_content:
                if any(kw in text.lower() for kw in keywords):
                    matches.append(text.strip())
            if matches:
                mappings.append({
                    "dimensao": dim,
                    "evidencias_encontradas": matches[:5],
                    "frequencia": len(matches),
                    "camada": 1 if dim in self.dimensions_c1 else 2
                })
        return mappings
    
    def parse_file(self, file_path, sheet_name=None):
        if file_path.endswith(".docx"): extracted = self.extract_from_docx(file_path)
        elif file_path.endswith(".xlsx"): extracted = self.extract_from_xlsx(file_path, sheet_name)
        else: return {"error": "Formato não suportado"}
        
        if "error" in extracted: return extracted
        
        mappings = self.map_to_dimensions(extracted)
        return {
            "source": extracted.get("source"),
            "type": extracted.get("type"),
            "mappings": mappings
        }